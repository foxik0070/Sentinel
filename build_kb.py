#!/usr/bin/env python3
"""
build_kb.py — Sentinel Commander Knowledge Base Builder

Reads docs from one or more directories, extracts and cleans text,
chunks it intelligently with overlap, deduplicates, and writes a
<<<SENTINEL_ENTRY>>>-delimited file for the RAG engine.

Usage:
    python3 build_kb.py
    python3 build_kb.py -s ./docs ./admindocs
    python3 build_kb.py -o /tmp/test.txt --dry-run
    python3 build_kb.py --no-meta -v
"""

import argparse
import csv
import subprocess
import hashlib
import json
import os
import re
import sys
import time
from pathlib import Path

# ─── OPTIONAL DEPENDENCIES ────────────────────────────────────────────────────

try:
    import docx as _docx
except ImportError:
    _docx = None

try:
    import pypdf as _pypdf
except ImportError:
    _pypdf = None

try:
    import pandas as _pd
except ImportError:
    _pd = None

try:
    import yaml as _yaml
except ImportError:
    _yaml = None

# ─── CONSTANTS ────────────────────────────────────────────────────────────────

DELIMITER   = "<<<SENTINEL_ENTRY>>>"
MAX_CHUNK   = 2500   # max chars per chunk
OVERLAP     = 300    # chars of tail from previous chunk kept as context seed

SCRIPT_DIR  = Path(__file__).resolve().parent

DEFAULT_SOURCES = [
    str(SCRIPT_DIR / "docs"),
    str(SCRIPT_DIR / "admindocs"),
    str(SCRIPT_DIR / "learning_knowledge_base"),
]
DEFAULT_OUTPUT  = str(SCRIPT_DIR / "knowledge_base.txt")

SENTINEL_CONFIG = "/etc/sentinel/config.yaml"

SUPPORTED = {
    ".md", ".txt", ".rst",
    ".docx", ".pdf",
    ".xlsx", ".xls", ".csv",
    ".yaml", ".yml",
    ".conf", ".ini",
}

# ─── TEXT EXTRACTION ──────────────────────────────────────────────────────────

def _read(path):
    return Path(path).read_text(encoding="utf-8", errors="replace")

def extract_md(path):
    return _read(path)

def extract_txt(path):
    return _read(path)

def extract_rst(path):
    """Convert RST underline headings to ## style."""
    lines = _read(path).split('\n')
    out, i = [], 0
    char_to_level = {'=': '#', '-': '##', '~': '###', '^': '####', '"': '#####'}
    while i < len(lines):
        if (i + 1 < len(lines)
                and lines[i].strip()
                and lines[i + 1].strip()
                and len(set(lines[i + 1].strip())) == 1
                and lines[i + 1][0] in char_to_level
                and len(lines[i + 1]) >= len(lines[i].strip())):
            level = char_to_level[lines[i + 1][0]]
            out.append(f"{level} {lines[i].strip()}")
            i += 2
        else:
            out.append(lines[i])
            i += 1
    return '\n'.join(out)

def extract_docx(path):
    if not _docx:
        return ""
    try:
        doc = _docx.Document(path)
        parts = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style = para.style.name
            if style.startswith('Heading'):
                try:
                    lvl = int(style.split()[-1])
                except ValueError:
                    lvl = 2
                parts.append('#' * min(lvl, 6) + ' ' + para.text.strip())
            else:
                parts.append(para.text.strip())
        for tbl in doc.tables:
            if not tbl.rows:
                continue
            headers = [c.text.strip() for c in tbl.rows[0].cells]
            parts.append("| " + " | ".join(headers) + " |")
            parts.append("| " + " | ".join(['---'] * len(headers)) + " |")
            for row in tbl.rows[1:]:
                parts.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")
        return '\n'.join(parts)
    except Exception as e:
        print(f"  [!] DOCX {Path(path).name}: {e}")
        return ""

def extract_pdf(path):
    if not _pypdf:
        return ""
    try:
        reader = _pypdf.PdfReader(path)
        parts = []
        for i, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(f"--- Page {i + 1} ---\n{text}")
        return '\n\n'.join(parts)
    except Exception as e:
        print(f"  [!] PDF {Path(path).name}: {e}")
        return ""

def extract_xlsx(path):
    if not _pd:
        return ""
    try:
        xf = _pd.ExcelFile(path)
        parts = []
        for sheet in xf.sheet_names:
            df = _pd.read_excel(xf, sheet_name=sheet, dtype=str).fillna('')
            if df.empty:
                continue
            parts.append(f"## Sheet: {sheet}")
            try:
                parts.append(df.to_markdown(index=False))
            except Exception:
                parts.append(df.to_string(index=False))
        return '\n\n'.join(parts)
    except Exception as e:
        print(f"  [!] XLSX {Path(path).name}: {e}")
        return ""

def extract_csv(path):
    if _pd:
        try:
            df = _pd.read_csv(path, dtype=str).fillna('')
            try:
                return df.to_markdown(index=False)
            except Exception:
                return df.to_string(index=False)
        except Exception as e:
            print(f"  [!] CSV {Path(path).name}: {e}")
            return ""
    try:
        rows = list(csv.reader(open(path, encoding='utf-8', errors='replace')))
        if not rows:
            return ""
        lines = ["| " + " | ".join(rows[0]) + " |",
                 "| " + " | ".join(['---'] * len(rows[0])) + " |"]
        lines += ["| " + " | ".join(r) + " |" for r in rows[1:]]
        return '\n'.join(lines)
    except Exception as e:
        print(f"  [!] CSV {Path(path).name}: {e}")
        return ""

def _redact(d, sensitive=('password', 'bind_password', 'token', 'secret_key',
                          'pass', 'api_key', 'jwt')):
    if isinstance(d, dict):
        for k in list(d.keys()):
            if any(s in k.lower() for s in sensitive):
                d[k] = "***REDACTED***"
            else:
                _redact(d[k], sensitive)
    elif isinstance(d, list):
        for item in d:
            _redact(item, sensitive)

def extract_yaml(path):
    raw = _read(path)
    if _yaml:
        try:
            data = _yaml.safe_load(raw)
            _redact(data)
            return "```yaml\n" + _yaml.dump(data, allow_unicode=True, default_flow_style=False) + "```"
        except Exception:
            pass
    # Fallback: redact inline
    return re.sub(
        r'(?im)^(\s*(?:password|bind_password|token|secret_key|pass|api_key)\s*:\s*).*$',
        r'\1***REDACTED***',
        raw
    )

def extract_conf(path):
    raw = _read(path)
    return re.sub(
        r'(?im)^(\s*(?:password|pass|token|secret)\s*[=:]\s*).*$',
        r'\1***REDACTED***',
        raw
    )

_EXTRACTOR_MAP = {
    ".md":   extract_md,
    ".txt":  extract_txt,
    ".rst":  extract_rst,
    ".docx": extract_docx,
    ".pdf":  extract_pdf,
    ".xlsx": extract_xlsx,
    ".xls":  extract_xlsx,
    ".csv":  extract_csv,
    ".yaml": extract_yaml,
    ".yml":  extract_yaml,
    ".conf": extract_conf,
    ".ini":  extract_conf,
}

def get_extractor(path):
    return _EXTRACTOR_MAP.get(Path(path).suffix.lower())

# ─── CHUNKER ──────────────────────────────────────────────────────────────────

_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$')

def chunk_text(text, max_size=MAX_CHUNK, overlap=OVERLAP):
    """
    Split text into (context_path, content) chunks.
    - Splits at heading boundaries to keep sections coherent.
    - Sections larger than max_size are split by paragraph with overlap.
    - Overlap: last `overlap` chars of a chunk are prepended to the next
      so the model has cross-boundary context.
    """
    lines = text.split('\n')
    sections = []       # [(breadcrumb_str, body_str)]
    breadcrumb = []
    cur = []

    for line in lines:
        m = _HEADING_RE.match(line)
        if m:
            body = '\n'.join(cur).strip()
            if body and len(body) > 20:
                sections.append((' > '.join(breadcrumb) or 'General', body))
            cur = []
            level = len(m.group(1))
            breadcrumb = breadcrumb[:level - 1] + [m.group(2).strip()]
        else:
            cur.append(line)

    body = '\n'.join(cur).strip()
    if body and len(body) > 20:
        sections.append((' > '.join(breadcrumb) or 'General', body))

    if not sections:
        stripped = text.strip()
        if stripped:
            sections = [('General', stripped)]

    chunks = []
    for ctx, body in sections:
        if len(body) <= max_size:
            chunks.append((ctx, body))
            continue

        # Split oversized section by paragraph with overlap
        paras = re.split(r'\n{2,}', body)
        buf, buf_len, part, tail = [], 0, 1, ""
        for para in paras:
            if buf_len + len(para) > max_size and buf:
                content = (tail + '\n\n' if tail else '') + '\n\n'.join(buf)
                chunks.append((f"{ctx} (Part {part})", content.strip()))
                tail = content[-overlap:] if len(content) > overlap else content
                part += 1
                buf, buf_len = [], 0
            buf.append(para)
            buf_len += len(para)
        if buf:
            content = (tail + '\n\n' if tail else '') + '\n\n'.join(buf)
            chunks.append((f"{ctx} (Part {part})", content.strip()))

    return chunks

# ─── SENTINEL META ────────────────────────────────────────────────────────────

def _sentinel_meta():
    """Auto-generates a KB section from live Sentinel system info."""
    parts = ["# Sentinel Commander — System Overview\n"]

    # Config (sanitized)
    try:
        raw = Path(SENTINEL_CONFIG).read_text(encoding="utf-8", errors="replace")
        redacted = re.sub(
            r'(?im)^(\s*(?:password|bind_password|token|secret_key|pass|api_key)\s*:\s*).*$',
            r'\1***REDACTED***',
            raw
        )
        parts.append("## Active Configuration\n```yaml\n" + redacted.strip() + "\n```\n")
    except Exception:
        pass

    # Plugins
    plugin_dir = SCRIPT_DIR / "sentinel" / "plugins"
    if plugin_dir.exists():
        plugins = sorted(p.name for p in plugin_dir.glob("*.py") if not p.name.startswith('_'))
        if plugins:
            parts.append("## Installed Detector Plugins\n")
            parts += [f"- {p}" for p in plugins]
            parts.append("")

    # Docs index (give the model an idea of what docs exist)
    docs_dir = SCRIPT_DIR / "docs"
    if docs_dir.exists():
        docs = sorted(p.name for p in docs_dir.glob("*") if p.is_file())
        if docs:
            parts.append("## Available Documentation Files\n")
            parts += [f"- {d}" for d in docs]
            parts.append("")

    return '\n'.join(parts)

# ─── BUILDER ──────────────────────────────────────────────────────────────────

def clean(text):
    text = re.sub(r'!\[([^\]]*)\]\([^)]*\)', r'[Image: \1]', text)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    return text.strip()

def build_knowledge_base(source_dirs, output_file, include_meta=True,
                         dry_run=False, verbose=False):
    t0 = time.time()
    col_w = 55

    print(f"\n{'═'*60}")
    print(f"  Sentinel KB Builder")
    print(f"  Output : {output_file}")
    print(f"  Sources: {source_dirs}")
    print(f"{'═'*60}\n")

    all_chunks = []         # [(file_label, context, content)]
    seen = set()            # MD5 deduplication

    file_ok = 0
    file_skip = 0

    # ── Sentinel meta section ────────────────────────────────────────────────
    if include_meta:
        meta = _sentinel_meta()
        if meta.strip():
            for ctx, body in chunk_text(meta):
                h = hashlib.md5(body.encode()).hexdigest()
                if h not in seen:
                    seen.add(h)
                    all_chunks.append(("[SENTINEL-META]", ctx, body))
            print(f"  [META] Sentinel system section added "
                  f"({sum(1 for f,_,_ in all_chunks if f=='[SENTINEL-META]')} chunks)\n")

    # ── Source directories ────────────────────────────────────────────────────
    for source_dir in source_dirs:
        src = Path(source_dir)
        if not src.exists():
            if verbose:
                print(f"  [SKIP] Not found: {source_dir}")
            continue

        print(f"  Scanning: {source_dir}")
        dir_chunks = 0

        for filepath in sorted(src.rglob('*')):
            if not filepath.is_file():
                continue
            # Skip hidden files, caches, binaries
            if any(part.startswith('.') or part in ('__pycache__', 'node_modules')
                   for part in filepath.parts):
                continue
            if filepath.suffix.lower() not in SUPPORTED:
                continue

            extractor = get_extractor(filepath)
            if not extractor:
                continue

            rel = filepath.relative_to(src)

            try:
                raw = extractor(filepath)
            except Exception as e:
                print(f"  [ERR] {rel}: {e}")
                file_skip += 1
                continue

            if not raw or not raw.strip():
                file_skip += 1
                if verbose:
                    print(f"  [EMPTY] {rel}")
                continue

            cleaned = clean(raw)
            file_new = 0

            for ctx, body in chunk_text(cleaned):
                if len(body) < 30:
                    continue
                h = hashlib.md5(body.encode()).hexdigest()
                if h in seen:
                    continue
                seen.add(h)
                all_chunks.append((str(rel), ctx, body))
                file_new += 1

            if file_new:
                file_ok += 1
                dir_chunks += file_new
                label = str(rel)[:col_w].ljust(col_w)
                print(f"    + {label}  {file_new:>3} chunks")
            else:
                file_skip += 1
                if verbose:
                    print(f"  [DUP ] {rel}")

        print(f"  → {dir_chunks} chunks from this directory\n")

    # ── Summary ──────────────────────────────────────────────────────────────
    total_chars = sum(len(c) for _, _, c in all_chunks)
    elapsed = time.time() - t0

    print(f"{'─'*60}")
    print(f"  Files OK      : {file_ok}")
    print(f"  Files skipped : {file_skip}")
    print(f"  Total chunks  : {len(all_chunks)}")
    print(f"  Total chars   : {total_chars:,}")
    print(f"  Time          : {elapsed:.1f}s")

    if dry_run:
        print(f"\n  [DRY RUN] Nothing written.")
        print(f"{'═'*60}\n")
        return

    # ── Write ────────────────────────────────────────────────────────────────
    out = Path(output_file)
    out.parent.mkdir(parents=True, exist_ok=True)

    with out.open('w', encoding='utf-8') as f:
        for file_label, ctx, content in all_chunks:
            f.write(f"{DELIMITER}\n")
            f.write(f"FILE: {file_label}\n")
            f.write(f"CONTEXT: {ctx}\n")
            f.write(f"CONTENT:\n{content}\n\n")

    print(f"  Written       : {out}  ({out.stat().st_size:,} bytes)")
    print(f"{'═'*60}")
    print(f"\n  RAG will auto-index on next Sentinel startup (mtime changed).")
    print(f"  To force re-index now: sudo systemctl restart sentinel\n")


# ─── CLI ──────────────────────────────────────────────────────────────────────

def _check_deps():
    missing = []
    for name, pkg in [("pypdf", "pypdf"), ("docx", "python-docx"),
                      ("pandas", "pandas"), ("yaml", "pyyaml")]:
        try:
            __import__(name)
        except ImportError:
            missing.append(pkg)
    if missing:
        print(f"\n  ⚠  Volitelné balíčky nejsou nainstalovány (některé formáty budou přeskočeny):")
        print(f"     pip3 install {' '.join(missing)}\n")
    return missing


# ─── INTERAKTIVNÍ MENU ────────────────────────────────────────────────────────

def _color(text, code):
    return f"\033[{code}m{text}\033[0m"

def _bold(t):   return _color(t, "1")
def _cyan(t):   return _color(t, "96")
def _green(t):  return _color(t, "92")
def _yellow(t): return _color(t, "93")
def _red(t):    return _color(t, "91")
def _dim(t):    return _color(t, "2")


def _interactive_menu():
    """Interaktivní průvodce pro sestavení KB."""
    print("\n" + "═" * 60)
    print(_bold(_cyan("  🛡  Sentinel Commander — Knowledge Base Builder")))
    print("═" * 60)
    print(f"  Výstup:  {_yellow(DEFAULT_OUTPUT)}")
    print(f"  Zdroje:  {', '.join(DEFAULT_SOURCES)}")
    print("═" * 60)

    # Zkontroluj dostupné zdroje
    sources_found = []
    for s in DEFAULT_SOURCES:
        p = Path(s)
        if p.exists():
            count = sum(1 for f in p.rglob("*") if f.is_file() and f.suffix.lower() in SUPPORTED)
            sources_found.append((s, count))
        else:
            sources_found.append((s, None))

    print("\n  Nalezené zdrojové adresáře:")
    for path, count in sources_found:
        if count is None:
            print(f"    {_red('✗')} {path} {_dim('(neexistuje)')}")
        else:
            print(f"    {_green('✓')} {path} {_dim(f'({count} souborů)')}")

    print(f"\n  Podporované formáty: {_dim('.md .txt .pdf .docx .xlsx .csv .yaml .conf .rst')}")
    print()

    print("  " + "─" * 56)
    print(f"  {_bold('1')}  Sestavit KB (výchozí zdroje)")
    print(f"  {_bold('2')}  Sestavit KB + přidat vlastní adresář nebo soubory")
    print(f"  {_bold('3')}  Dry-run (zobrazit co bude zahrnuto bez zápisu)")
    print(f"  {_bold('4')}  Instalovat chybějící závislosti")
    print(f"  {_bold('5')}  Zobrazit obsah aktuální KB")
    print(f"  {_bold('q')}  Konec")
    print("  " + "─" * 56)

    while True:
        try:
            choice = input(f"\n  {_cyan('Volba')} [1-5/q]: ").strip().lower()
        except (KeyboardInterrupt, EOFError):
            print("\n  Přerušeno.")
            return

        if choice in ("q", "quit", "exit"):
            print("  Na shledanou!")
            return

        elif choice in ("1", ""):
            missing = _check_deps()
            build_knowledge_base(
                source_dirs=DEFAULT_SOURCES,
                output_file=DEFAULT_OUTPUT,
                include_meta=True,
                dry_run=False,
                verbose=True,
            )
            _kb_stats(DEFAULT_OUTPUT)
            return

        elif choice == "2":
            print(f"\n  Zadej cestu k adresáři nebo souboru(ům):")
            print(f"  {_dim('(Enter = hotovo, prázdné = konec)')}")
            extra = []
            while True:
                try:
                    raw = input(f"  Přidat: ").strip()
                except (KeyboardInterrupt, EOFError):
                    break
                if not raw:
                    break
                p = Path(raw)
                if not p.exists():
                    print(f"  {_red('✗')} Nenalezeno: {raw}")
                else:
                    extra.append(raw)
                    print(f"  {_green('✓')} Přidáno: {raw}")
            sources = DEFAULT_SOURCES + extra
            missing = _check_deps()
            build_knowledge_base(
                source_dirs=sources,
                output_file=DEFAULT_OUTPUT,
                include_meta=True,
                dry_run=False,
                verbose=True,
            )
            _kb_stats(DEFAULT_OUTPUT)
            return

        elif choice == "3":
            _check_deps()
            build_knowledge_base(
                source_dirs=DEFAULT_SOURCES,
                output_file=DEFAULT_OUTPUT,
                include_meta=True,
                dry_run=True,
                verbose=True,
            )
            return

        elif choice == "4":
            _install_deps()
            input(f"\n  {_dim('Stiskni Enter pro pokračování...')}")

        elif choice == "5":
            _show_kb_info(DEFAULT_OUTPUT)
            input(f"\n  {_dim('Stiskni Enter pro pokračování...')}")

        else:
            print(f"  {_yellow('?')} Neznámá volba. Zadej 1-5 nebo q.")


def _kb_stats(output_file):
    """Zobrazí stručné statistiky o vygenerované KB."""
    p = Path(output_file)
    if not p.exists():
        return
    content = p.read_text(encoding="utf-8", errors="replace")
    entries = content.count("<<<SENTINEL_ENTRY>>>")
    size_kb = p.stat().st_size // 1024
    print(f"\n  {_green('✓')} KB připravena:")
    print(f"    Soubor:  {output_file}")
    print(f"    Velikost: {size_kb} KB")
    print(f"    Chunků:  {entries}")
    print(f"\n  {_dim('Sentinel automaticky reindexuje KB při příštím startu.')}")
    print(f"  {_dim('Pro okamžitý reindex: sudo systemctl restart sentinel')}")
    print(f"  {_dim('nebo v UI: Tools → KB → Reindex')}")


def _show_kb_info(output_file):
    """Zobrazí info o aktuální KB."""
    p = Path(output_file)
    if not p.exists():
        print(f"\n  {_red('✗')} KB soubor neexistuje: {output_file}")
        return
    content = p.read_text(encoding="utf-8", errors="replace")
    entries = content.count("<<<SENTINEL_ENTRY>>>")
    size_kb = p.stat().st_size // 1024
    mtime = time.strftime("%Y-%m-%d %H:%M", time.localtime(p.stat().st_mtime))
    print(f"\n  Aktuální knowledge_base.txt:")
    print(f"    Cesta:     {output_file}")
    print(f"    Velikost:  {size_kb} KB  ({p.stat().st_size:,} bytů)")
    print(f"    Chunků:    {entries}")
    print(f"    Upraveno:  {mtime}")
    # Prvních 5 chunků
    chunks = content.split("<<<SENTINEL_ENTRY>>>")[1:6]
    if chunks:
        print(f"\n  Prvních {len(chunks)} chunků (preview):")
        for i, ch in enumerate(chunks, 1):
            preview = ch.strip()[:120].replace("\n", " ")
            print(f"    [{i}] {_dim(preview)}")


def _install_deps():
    """Nainstaluje chybějící volitelné závislosti."""
    deps = [("pypdf", "pypdf"), ("docx", "python-docx"),
            ("pandas", "pandas"), ("yaml", "pyyaml")]
    for name, pkg in deps:
        try:
            __import__(name)
            print(f"  {_green('✓')} {pkg} — již nainstalováno")
        except ImportError:
            print(f"  Instaluji {pkg}...")
            ret = subprocess.run(
                [sys.executable, "-m", "pip", "install", pkg],
                capture_output=True, text=True
            )
            if ret.returncode == 0:
                print(f"  {_green('✓')} {pkg} nainstalováno")
            else:
                print(f"  {_red('✗')} {pkg}: {ret.stderr.strip()[:100]}")


def main():
    p = argparse.ArgumentParser(
        description="Sentinel Commander — Knowledge Base Builder",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Spuštění bez argumentů → interaktivní menu.

Příklady:
  python3 build_kb.py                         # interaktivní menu
  python3 build_kb.py --auto                  # přímé sestavení bez menu
  python3 build_kb.py -s ./docs ./admindocs   # vlastní zdroje
  python3 build_kb.py -o /tmp/test.txt --dry-run
  python3 build_kb.py --no-meta -v

Podporované formáty:
  Markdown (.md), Text (.txt), reStructuredText (.rst),
  PDF (.pdf), Word (.docx), Excel (.xlsx/.xls),
  CSV (.csv), YAML (.yaml/.yml), Config (.conf/.ini)
        """
    )
    p.add_argument("-s", "--sources", nargs="+", default=None,
                   metavar="DIR",
                   help=f"Zdrojové adresáře (výchozí: {DEFAULT_SOURCES})")
    p.add_argument("-o", "--output", default=DEFAULT_OUTPUT,
                   help=f"Výstupní soubor (výchozí: {DEFAULT_OUTPUT})")
    p.add_argument("--no-meta", action="store_true",
                   help="Přeskočit auto-generovanou Sentinel meta sekci")
    p.add_argument("--dry-run", action="store_true",
                   help="Zobrazit co bude zahrnuto bez zápisu souboru")
    p.add_argument("-v", "--verbose", action="store_true",
                   help="Zobrazit i přeskočené/duplicitní soubory")
    p.add_argument("--auto", action="store_true",
                   help="Sestavit bez interaktivního menu (výchozí zdroje)")
    args = p.parse_args()

    # Interaktivní menu pokud nejsou argumenty
    if not args.auto and args.sources is None and not args.dry_run and not args.no_meta:
        _interactive_menu()
        return

    _check_deps()
    build_knowledge_base(
        source_dirs=args.sources or DEFAULT_SOURCES,
        output_file=args.output,
        include_meta=not args.no_meta,
        dry_run=args.dry_run,
        verbose=args.verbose,
    )
    if not args.dry_run:
        _kb_stats(args.output)

if __name__ == "__main__":
    main()
